from pathlib import Path

from docx import Document


class SongInstance(object):
    def __init__(self, song, artist, language):
        self.artist = artist.capitalize()
        self.song = song.capitalize()
        self.language = language
        self.text = None
        self.translation = None

    def set_text(self, text):
        self.text = text

    def set_translation(self, translation):
        self.translation = translation

    def _get_header(self):
        return f"{self.song} - {self.artist} in {self.language}"

    def _get_name(self):
        return f"{self.language}__{self.song}__{self.artist}".replace(" ", "_").lower()

    def save(self, folder, kind="word"):

        self.header = self._get_header()
        self.name = self._get_name()

        if kind == "word":
            document = Document()
            document.add_heading(self.header, level=1)
            document.add_paragraph(self.text)
            document.save(Path(folder / f"{self.name}.docx"))

    def __str__(self):
        seperator = "----" * 10
        string = [self.text, seperator, self.translation]
        return "\n".join(string)
