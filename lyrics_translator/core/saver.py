from pathlib import Path

from docx import Document


class Saver(object):
    def __init__(self, song: str, artist: str, translation: str, language: str):
        self.song = song
        self.artist = artist
        self.translation = translation
        self.language = language

    def _get_header(self) -> str:
        """_summary_

        Returns:
            str: _description_
        """
        return f"{self.song} - {self.artist} in {self.language}"

    def _get_name(self) -> str:
        """_summary_

        Returns:
            str: _description_
        """
        return f"{self.language}__{self.song}__{self.artist}".replace(" ", "_").lower()

    def to_word(self, folder: Path):
        document = Document()
        document.add_heading(self.header, level=1)
        document.add_paragraph(self.translation)
        document.save(Path(folder / f"{self.name}.docx"))

    def to_txt(self, folder: Path):
        with open(Path(folder / f"{self.name}.txt"), mode="w", encoding="utf-8") as f:
            f.write(self.translation)

    def save(self, folder: Path, kind: str = "txt") -> None:
        """_summary_

        Args:
            folder (Path): _description_
            kind (str, optional): _description_. Defaults to "txt".
        """
        self.header = self._get_header()
        self.name = self._get_name()
        if kind == "word":
            self.to_word(folder)
        elif kind == "txt":
            self.to_txt(folder)
