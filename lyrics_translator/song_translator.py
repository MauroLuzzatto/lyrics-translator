import warnings
from pathlib import Path
from typing import List, Tuple

import lyricsgenius
from docx import Document
from dotenv import dotenv_values

from lyrics_translator.translator import Translator

config = dotenv_values(".env")
genius = lyricsgenius.Genius(config["GENIUS_ACCESS_TOKEN"])


class SongTranslator(object):
    def __init__(self, song: str, artist: str, language: str = "de"):
        self.song = song
        self.artist = artist
        self.language = language
        self.text = None
        self.translation = None

        self.translator = Translator(language)
        self.translator.get_translator_pipeline()

    def get_song_translations(self):
        self.download_lyrics()
        self.translate()

    def download_lyrics(self, get_full_info=False):

        try:
            genius_song = genius.search_song(
                self.song, self.artist, get_full_info=get_full_info
            )
        except KeyError:
            pass

        if not genius_song:
            message = f"not found - song: {self.song}, artist: {self.artist}"
            warnings.warn(message)
            self.text = ""
        else:
            self.text = genius_song.lyrics
        return self.text

    def translate(self):
        self.translation = self.translator.translate(self.text)
        return self.translation

    def _get_header(self):
        return f"{self.song} - {self.artist} in {self.language}"

    def _get_name(self):
        return f"{self.language}__{self.song}__{self.artist}".replace(" ", "_").lower()

    def save(self, folder: Path, kind: str = "txt"):
        """_summary_

        Args:
            folder (Path): _description_
            kind (str, optional): _description_. Defaults to "txt".
        """

        self.header = self._get_header()
        self.name = self._get_name()

        if kind == "word":
            document = Document()
            document.add_heading(self.header, level=1)
            document.add_paragraph(self.translation)
            document.save(Path(folder / f"{self.name}.docx"))

        elif kind == "txt":
            with open(Path(folder / f"{self.name}.txt"), "w") as f:
                f.write(self.translation)

    def __str__(self):
        seperator = "----" * 10
        string = [self.text, seperator, self.translation]
        return "\n".join(string)
