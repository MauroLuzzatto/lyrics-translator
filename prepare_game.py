from pathlib import Path

from docx import Document
from dotenv import dotenv_values

from lyrics_translator import LyricsTranslator, SpotifyAPI, create_folder, get_base_path

if __name__ == "__main__":

    base_path = get_base_path()
    resources_path = Path(base_path / "resources")
    folder = create_folder(Path(base_path / "lyrics"))
    config = dotenv_values(".env")

    user_id = 1157239771
    playlist_id = "3VQvmc9cjrVj7HcjxwN5Jg"
    language = "sv"
    document_name = "Midsummer 2022"

    spotifAPI = SpotifyAPI(
        user_id, config["SPOTIFY_CLIENT_ID"], config["SPOTIFY_CLIENT_SECRET"]
    )
    playlist = spotifAPI.get_playlist_items(playlist_id=playlist_id)

    songs = []
    for track in playlist["items"]:
        artist = ", ".join([artist["name"] for artist in track["track"]["artists"]])
        songs.append((track["track"]["name"], artist, language))

    print(songs)

    document = Document()

    for index, (song, artist, language) in enumerate(songs):
        lyrics = LyricsTranslator(song, artist, config, language, testing=False)

        lyrics.get_song_translation()
        heading = lyrics.translator.translate(song)

        document.add_heading(heading, level=1)
        document.add_paragraph(lyrics.translation)

    document.save(Path(folder / f"{document_name}.docx"))
